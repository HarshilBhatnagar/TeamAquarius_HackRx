import requests
import pdfplumber
import docx
from typing import List, Dict, Any, Optional
from utils.logger import logger
import re
import io

import asyncio
import aiohttp

async def get_document_text(url: str) -> str:
    """
    ROUND 2 AGENTIC DOCUMENT EXTRACTION: Let the LLM handle all document types naturally.
    Handles diverse document structures including multi-column layouts, tables, and complex PDFs.
    """
    try:
        logger.info(f"ROUND 2 AGENTIC: Downloading document from: {url}")
        
        # Use aiohttp for async download
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=aiohttp.ClientTimeout(total=30)) as response:
                response.raise_for_status()
                content = await response.read()
        
        content_type = response.headers.get('content-type', '').lower()
        
        if 'pdf' in content_type or url.lower().endswith('.pdf'):
            text = extract_pdf_text(content)
            return text
        elif 'docx' in content_type or url.lower().endswith('.docx'):
            text = extract_docx_text(content)
            return text
        else:
            # Try to detect PDF by content
            if content.startswith(b'%PDF'):
                text = extract_pdf_text(content)
                return text
            else:
                raise ValueError(f"Unsupported document type: {content_type}")
                
    except Exception as e:
        logger.error(f"Error downloading document: {e}")
        raise

def extract_pdf_text(pdf_content: bytes) -> str:
    """
    HYBRID PDF EXTRACTION: Table-aware parsing with speed optimization.
    Combines best of both approaches for maximum accuracy and speed.
    """
    try:
        text_content = []
        
        with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"Processing PDF with {total_pages} pages")
            
            # Limit pages for performance - most documents don't exceed 25 pages
            max_pages = min(total_pages, 25)
            if total_pages > 25:
                logger.warning(f"Large document detected ({total_pages} pages). Limiting to first 25 pages for performance.")
            
            for page_num in range(max_pages):
                page = pdf.pages[page_num]
                logger.info(f"Processing page {page_num + 1}/{max_pages}")
                
                # HYBRID APPROACH: Table-aware + layout preservation
                page_text = extract_page_text_hybrid(page)
                text_content.append(page_text)
                
                # Early termination if document is getting too large
                current_size = sum(len(text) for text in text_content)
                if current_size > 500000:  # 500KB limit
                    logger.warning(f"Document size limit reached ({current_size} chars). Stopping at page {page_num + 1}")
                    break
        
        full_text = "\n\n".join(text_content)
        logger.info(f"Extracted {len(full_text)} characters from PDF (processed {len(text_content)} pages)")
        
        # CRITICAL FIX: Validate document content
        if len(full_text) > 500000:
            logger.warning(f"Document too large ({len(full_text)} chars), may be wrong document")
            # Check if it's the wrong document by looking for specific keywords
            if "arogya sanjeevani" in full_text.lower() and "hdfc" not in full_text.lower():
                logger.error("Wrong document detected: Arogya Sanjeevani instead of HDFC Life Insurance")
                raise ValueError("Wrong document detected - processing Arogya Sanjeevani instead of HDFC Life Insurance")
        
        return full_text
        
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        raise

def extract_page_text_hybrid(page) -> str:
    """
    HYBRID APPROACH: Table-aware parsing with layout preservation.
    Combines best of both approaches for maximum accuracy.
    """
    page_content = []
    
    try:
        # 1. Extract tables first with enhanced formatting (from suggested approach)
        tables = page.extract_tables()
        if tables:
            logger.info(f"Found {len(tables)} tables on page")
            for table_idx, table in enumerate(tables):
                if table:
                    # Use enhanced table formatting
                    markdown_table = format_table_enhanced(table)
                    page_content.append(f"\n\n--- TABLE {table_idx + 1} START ---\n{markdown_table}\n--- TABLE {table_idx + 1} END ---\n\n")
        
        # 2. Extract text with layout awareness (our approach)
        try:
            # Use chars method which is more reliable across pdfplumber versions
            chars = page.chars
            if chars:
                # Group characters by lines and columns
                lines = {}
                for char in chars:
                    y_pos = round(char['y0'], 2)
                    if y_pos not in lines:
                        lines[y_pos] = []
                    lines[y_pos].append(char)
                
                # Sort lines by y position and characters by x position
                sorted_lines = sorted(lines.items())
                for y_pos, line_chars in sorted_lines:
                    line_chars.sort(key=lambda c: c['x0'])
                    line_text = ''.join([c['text'] for c in line_chars])
                    if line_text.strip():
                        page_content.append(line_text)
        except Exception as layout_error:
            logger.warning(f"Error in layout-aware extraction: {layout_error}, falling back to plain text")
            # Don't add anything to page_content, let it fall through to plain text
        
        # 3. Fallback: Extract plain text if layout extraction fails
        if not page_content:
            plain_text = page.extract_text()
            if plain_text:
                page_content.append(plain_text)
        
        return "\n\n".join(page_content)
        
    except Exception as e:
        logger.warning(f"Error in hybrid extraction: {e}, falling back to plain text")
        return page.extract_text() or ""

def format_table_enhanced(table_data):
    """
    Enhanced table formatting from suggested approach.
    Converts table data to clean markdown format.
    """
    if not table_data:
        return ""
    
    try:
        # Clean table data
        cleaned_table = [[(str(cell) or "").replace("\n", " ").strip() for cell in row] for row in table_data]
        
        # Create markdown table
        header = " | ".join(cleaned_table[0])
        markdown_table = f"| {header} |\n"
        
        # Add separator
        separator = " | ".join(["---"] * len(cleaned_table[0]))
        markdown_table += f"| {separator} |\n"
        
        # Add data rows
        for row in cleaned_table[1:]:
            # Ensure row has same number of columns as header
            while len(row) < len(cleaned_table[0]):
                row.append("")
            data_row = " | ".join(row)
            markdown_table += f"| {data_row} |\n"
        
        return markdown_table
        
    except Exception as e:
        logger.warning(f"Error formatting table: {e}")
        return str(table_data)

def process_table(table: List[List[str]]) -> str:
    """
    Process tables with enhanced formatting.
    Handles dense tables of information and policy details.
    """
    if not table:
        return ""
    
    try:
        table_text = []
        
        # Clean and format table data
        for row_idx, row in enumerate(table):
            if not row:
                continue
                
            # Clean row data
            cleaned_row = []
            for cell in row:
                if cell:
                    # Clean cell text
                    cell_text = str(cell).strip()
                    # Remove excessive whitespace
                    cell_text = re.sub(r'\s+', ' ', cell_text)
                    cleaned_row.append(cell_text)
                else:
                    cleaned_row.append("")
            
            # Skip empty rows
            if not any(cell for cell in cleaned_row):
                continue
            
            # Format row based on content
            if row_idx == 0:
                # Header row
                table_text.append(" | ".join(cleaned_row))
                table_text.append("-" * 50)  # Separator
            else:
                # Data row
                row_text = " | ".join(cleaned_row)
                table_text.append(row_text)
        
        return "\n".join(table_text)
        
    except Exception as e:
        logger.warning(f"Error processing table: {e}")
        return ""

def extract_docx_text(docx_content: bytes) -> str:
    """
    Extract text from DOCX documents.
    """
    try:
        doc = docx.Document(io.BytesIO(docx_content))
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract tables from DOCX
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))
            
            if table_text:
                text_content.append("TABLE:\n" + "\n".join(table_text))
        
        full_text = "\n\n".join(text_content)
        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        return full_text
        
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        raise